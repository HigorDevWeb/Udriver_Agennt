from flask import Flask, request, jsonify
import openai
import os
from dotenv import load_dotenv
import tempfile
from flask_cors import CORS
import requests
from docx import Document

# Load variables from .env file
load_dotenv()
openai.api_key = os.getenv("OPEN_AI_API_KEY")
DIFY_API_KEY = os.getenv("DIFY_API_KEY")
DIFY_API_URL = os.getenv("DIFY_API_URL")
DIFY_DATASET_ID = os.getenv("DIFY_DATASET_ID")

app = Flask(__name__)
CORS(app)

def translate_text(text, target_language):
    import openai
    client = openai.OpenAI(api_key=os.getenv("OPEN_AI_API_KEY"))
    language_names = {
        "en": "English",
        "pl": "Polish",
        "uk": "Ukrainian",
        "ru": "Russian"
    }[target_language]
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"Translate to {language_names} in a natural and professional way."},
            {"role": "user", "content": text}
        ],
        max_tokens=2048,
        temperature=0.1,
    )
    return response.choices[0].message.content.strip()

def generate_docx(translations, base_name):
    document = Document()
    document.add_heading('Multilingual Transcription', 0)
    for language, content in translations.items():
        document.add_heading(language, level=1)
        document.add_paragraph(content)
    filename = f"{base_name}_multilingual.docx"
    document.save(filename)
    return filename

def send_to_dify(docx_filename):
    url = f"{DIFY_API_URL}/v1/datasets/{DIFY_DATASET_ID}/document/create-by-file"
    headers = {'Authorization': f'Bearer {DIFY_API_KEY}'}
    data = {'data': '{"indexing_technique":"high_quality","process_rule":{"mode":"automatic"}}'}
    
    print(f"Sending to Dify URL: {url}")
    print(f"Using API Key: {DIFY_API_KEY[:10]}...")
    print(f"Dataset ID: {DIFY_DATASET_ID}")
    
    try:
        with open(docx_filename, 'rb') as file_obj:
            files = {
                'file': (docx_filename, file_obj, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            response = requests.post(url, headers=headers, data=data, files=files)
        return response.status_code, response.text
    except Exception as e:
        print(f"Error sending to Dify: {e}")
        return 500, str(e)

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    if 'file' not in request.files:
        return jsonify({'error': 'No file in request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as temp:
        file.save(temp.name)
        file_path = temp.name

    try:
        with open(file_path, "rb") as audio_file:
            transcription = openai.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )
        text = transcription.text

        # Translation to 4 languages
        translations = {
            "Portuguese": text,
            "English": translate_text(text, "en"),
            "Polish": translate_text(text, "pl"),
            "Ukrainian": translate_text(text, "uk"),
            "Russian": translate_text(text, "ru"),
        }

        # Generate DOCX document
        base_name = os.path.splitext(file.filename)[0]
        docx_filename = generate_docx(translations, base_name)

        # Send document to Dify
        status, response = send_to_dify(docx_filename)
        print(f"Dify Response Status: {status}")
        print(f"Dify Response: {response}")
        
        # Clean up the docx file after sending
        try:
            os.remove(docx_filename)
        except Exception as cleanup_error:
            print(f"Warning: Could not remove docx file: {cleanup_error}")

        return jsonify({
            'transcription': text,
            'translations': translations,
            'docx_sent': status == 200,
            'dify_response': response,
            'dify_status': status
        })

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500
    finally:
        os.remove(file_path)

if __name__ == '__main__':
    app.run(debug=True)
